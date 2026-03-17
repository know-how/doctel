from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def _set_run_font(run, *, size_pt: int | None = None, bold: bool | None = None) -> None:
    font = run.font
    font.name = "Calibri"
    if size_pt is not None:
        font.size = Pt(size_pt)
    if bold is not None:
        font.bold = bold


def _add_title_page(doc: Document, *, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    _set_run_font(run, size_pt=26, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    _set_run_font(run, size_pt=14, bold=False)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Date: {date.today().isoformat()}")
    _set_run_font(run, size_pt=11, bold=False)

    doc.add_page_break()


def _add_toc_placeholder(doc: Document) -> None:
    doc.add_heading("Table of Contents", level=1)
    p = doc.add_paragraph(
        "Update this table in Microsoft Word: References → Table of Contents → Update Table."
    )
    _set_run_font(p.runs[0], size_pt=11, bold=False)
    doc.add_paragraph("1.0 ABSTRACT")
    doc.add_paragraph("2.0 INTRODUCTION")
    doc.add_paragraph("3.0 PROBLEM DEFINITION")
    doc.add_paragraph("4.0 OVERVIEW OF THE DEVELOPED APPLICATION")
    doc.add_paragraph("5.0 OBJECTIVES")
    doc.add_paragraph("6.0 USE CASE DIAGRAM")
    doc.add_paragraph("7.0 SCREENSHOTS")
    doc.add_paragraph("8.0 CONCLUSION")
    doc.add_page_break()


def _add_section(doc: Document, heading: str, body_paragraphs: list[str]) -> None:
    doc.add_heading(heading, level=1)
    for text in body_paragraphs:
        if text.strip().startswith("- "):
            p = doc.add_paragraph(text.strip()[2:], style="List Bullet")
            _set_run_font(p.runs[0], size_pt=11, bold=False)
        else:
            p = doc.add_paragraph(text)
            if p.runs:
                _set_run_font(p.runs[0], size_pt=11, bold=False)


def _add_use_case_diagram(doc: Document) -> None:
    doc.add_heading("6.0 USE CASE DIAGRAM", level=1)
    doc.add_paragraph(
        "Below is a text-based use case diagram representation (UML-style) for DocIntel."
    )
    diagram = (
        "+----------------------+\\n"
        "|        User          |\\n"
        "+----------+-----------+\\n"
        "           |\\n"
        "           v\\n"
        "+----------------------+\\n"
        "|     Login/Session    |\\n"
        "+----------+-----------+\\n"
        "           |\\n"
        "           +-------------------------------+\\n"
        "           |                               |\\n"
        "           v                               v\\n"
        "+--------------------------+     +---------------------------+\\n"
        "| Upload Document          |     | View Document Analysis    |\\n"
        "| - Choose PDF/DOCX/TXT    |     | - Executive Summary       |\\n"
        "| - Enter Metadata         |     | - Detailed Summary        |\\n"
        "| - Submit Upload          |     | - Topics/Entities         |\\n"
        "+------------+-------------+     | - Sentiment               |\\n"
        "             |                   | - Action items/Decisions  |\\n"
        "             v                   +-------------+-------------+\\n"
        "+--------------------------+                   |\\n"
        "| Store Document Securely  |                   v\\n"
        "+--------------------------+     +---------------------------+\\n"
        "                              -> | View Suggested Prompts    |\\n"
        "                                 | - 3–5 prompt suggestions  |\\n"
        "                                 +-------------+-------------+\\n"
        "                                               |\\n"
        "                                               v\\n"
        "                                 +---------------------------+\\n"
        "                                 | Ask Question (Chat)       |\\n"
        "                                 | - Select prompt or type   |\\n"
        "                                 | - Receive grounded answer |\\n"
        "                                 | - See cited snippets      |\\n"
        "                                 +-------------+-------------+\\n"
        "                                               |\\n"
        "                                               v\\n"
        "                                 +---------------------------+\\n"
        "                                 | View Original Document    |\\n"
        "                                 | - Open/download uploaded  |\\n"
        "                                 +---------------------------+\\n"
    )
    p = doc.add_paragraph(diagram)
    if p.runs:
        _set_run_font(p.runs[0], size_pt=10, bold=False)


def _add_screenshots_section(doc: Document) -> None:
    doc.add_heading("7.0 SCREENSHOTS", level=1)
    doc.add_paragraph(
        "Insert screenshots below. For each screenshot, replace the placeholder caption with the actual image and caption."
    )

    captions = [
        "Figure 1: Login Screen (DocIntel Portal)",
        "Figure 2: Document Upload Modal (metadata fields + file picker)",
        "Figure 3: Document View – Analysis Dashboard (Left Panel)",
        "Figure 4: Document View – Copilot (Right Panel)",
        "Figure 5: View Original Document (open/download)",
        "Figure 6 (Optional): Mobile App – Upload Screen",
        "Figure 7 (Optional): Mobile App – Chat Screen",
    ]

    for caption in captions:
        doc.add_paragraph("")
        p = doc.add_paragraph("[Insert screenshot here]")
        if p.runs:
            _set_run_font(p.runs[0], size_pt=11, bold=False)
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if p.runs:
            _set_run_font(p.runs[0], size_pt=10, bold=True)


def build_report_docx(output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    _add_title_page(
        doc,
        title="Document Analysis System (DocIntel)",
        subtitle="System Report (Based on the DocIntel Functional Specification Document)",
    )

    _add_toc_placeholder(doc)

    doc.add_heading("NAME OF APPLICATION: DocIntel (Document Analysis System)", level=1)

    _add_section(
        doc,
        "1.0 ABSTRACT",
        [
            "DocIntel is an internal document analysis and knowledge-access system designed to help staff quickly understand and use information contained in company documents such as policies, reports, technical specifications, and meeting minutes.",
            "The system ingests documents in common formats (PDF, DOCX, TXT), stores them in a centralized repository, and produces structured outputs including an executive summary, detailed summary, identified entities (people, dates, locations), thematic topics, sentiment, and extracted action items/decisions.",
            "A core capability is contextual prompt generation: the system automatically suggests 3–5 questions relevant to the uploaded document, enabling rapid exploration without reading the entire file.",
            "Users can also ask custom questions through a chat-style interface; answers are grounded strictly in the document content.",
        ],
    )

    _add_section(
        doc,
        "2.0 INTRODUCTION",
        [
            "Organizations accumulate large volumes of documents that contain critical operational knowledge—requirements, procedures, approvals, risks, and decisions.",
            "However, most documents are lengthy and unstructured, making retrieval of specific information time-consuming and inconsistent across users.",
            "DocIntel acts as an “AI Analyst” for company knowledge: when a user uploads a document, the system reads it and presents summaries, key insights, and actionable prompts that guide user interaction with the content.",
        ],
    )

    _add_section(
        doc,
        "3.0 PROBLEM DEFINITION",
        [
            "The main challenges addressed by DocIntel include:",
            "- Slow information retrieval from long documents",
            "- Inconsistent interpretation of key points across readers",
            "- Low reuse of knowledge contained in static files",
            "- Difficulty onboarding and referencing past decisions and policies",
            "Problem statement: There is a need for an internal system that can ingest common document formats, store them securely, automatically summarize and analyze content, generate contextual prompts, and answer user questions based strictly on uploaded document content.",
        ],
    )

    _add_section(
        doc,
        "4.0 OVERVIEW OF THE DEVELOPED APPLICATION",
        [
            "DocIntel is implemented as a three-layer system:",
            "- Presentation layer: React + TypeScript web UI (responsive two-panel layout)",
            "- Application layer: FastAPI backend exposing REST endpoints for ingestion, analysis, prompts, and Q&A",
            "- Data layer: Centralized document storage (current: disk-based; future: database/object storage)",
            "Core modules include document ingestion & metadata, text extraction (PDF/DOCX/TXT), summarization & analysis, prompt generation, and grounded Q&A chat with citations.",
        ],
    )

    _add_section(
        doc,
        "5.0 OBJECTIVES",
        [
            "General objective: Develop an internal document analysis system that accelerates document comprehension and improves information retrieval by generating summaries, insights, and contextual prompts.",
            "Specific objectives:",
            "- Enable uploading and secure storage of PDF, DOCX, and TXT documents with metadata",
            "- Generate executive and detailed summaries after processing",
            "- Identify key entities (people, dates, locations), topics/themes, and sentiment",
            "- Extract action items and decisions when present",
            "- Provide a two-panel UI: analysis dashboard and prompt/chat interface",
            "- Ensure question answering is grounded strictly in the uploaded document",
        ],
    )

    _add_use_case_diagram(doc)

    _add_screenshots_section(doc)

    _add_section(
        doc,
        "8.0 CONCLUSION",
        [
            "DocIntel addresses the operational challenge of extracting actionable knowledge from internal documents quickly and consistently.",
            "By implementing ingestion, summarization, analysis, and contextual prompt generation, the system helps users interact with documents efficiently and reduces time spent searching through large files.",
            "The modular architecture supports future enhancements such as persistent storage, improved retrieval, and stronger local model integration for fully offline inference.",
        ],
    )

    doc.save(str(output_path))


def main() -> None:
    project_root = Path(__file__).resolve().parents[1]
    out_dir = project_root / "docs"
    output_path = out_dir / "DocIntel_System_Report.docx"
    build_report_docx(output_path)
    print(str(output_path))


if __name__ == "__main__":
    main()
