import re
import os
import hashlib
import json
import logging
import asyncio
import time
import shutil
import uuid
from datetime import datetime, timezone
from typing import List, Optional
from pathlib import Path
from PIL import Image
import pytesseract
from PyPDF2 import PdfReader
import docx
from sqlalchemy import select, text
from sqlalchemy.ext.asyncio import AsyncSession

_tesseract_checked = False
_tesseract_available = False

def _check_tesseract() -> bool:
    global _tesseract_checked, _tesseract_available
    if not _tesseract_checked:
        _tesseract_checked = True
        _tesseract_available = shutil.which("tesseract") is not None
        if not _tesseract_available:
            logger = logging.getLogger(__name__)
            logger.warning("tesseract not found in PATH – OCR will be skipped")
    return _tesseract_available

from app.config import settings
import asyncio

# Track auto-training cooldown (only trigger once per batch)
_last_auto_train = 0
AUTO_TRAIN_COOLDOWN_SEC = 300

def _schedule_auto_training(project_id: int) -> None:
    """Schedule automatic model training after document ingestion.

    Takes a plain integer ``project_id`` rather than an ORM ``Document``
    object to avoid any risk of the fire-and-forget background task
    accessing a detached ORM instance (which would trigger a
    ``greenlet_spawn`` / ``MissingGreenlet`` error).

    Creates its own database session so the fire-and-forget task does not
    share the caller's ``AsyncSession``, which may be closed by the time
    the background task runs.
    """
    global _last_auto_train
    import time
    now = time.time()
    if now - _last_auto_train < AUTO_TRAIN_COOLDOWN_SEC:
        return
    _last_auto_train = now

    if project_id <= 0:
        return

    async def _do_train():
        try:
            from app.db.database import AsyncSessionLocal
            from app.services.multi_model_trainer import train_models_from_project

            async with AsyncSessionLocal() as train_db:
                result = await train_models_from_project(
                    project_ids=[project_id],
                    db=train_db,
                )
                import logging
                logging.getLogger().info(f"Auto-training complete for project {project_id}: {result}")
        except Exception as e:
            import logging
            logging.getLogger().warning(f"Auto-training skipped: {e}")

    asyncio.ensure_future(_do_train())
from app.db.models import Document, DocAnalysis, SuggestedPrompt, Chunk, Embedding
from app.utils.ollama_client import ollama
from app.services.model_resolver_service import resolve_model
from app.config import settings as app_settings
from app.utils.chroma_client import chroma
from app.utils.pgvector_client import insert_chunks as pgvector_insert_chunks
from app.services.embedding_service import (
    generate_embedding,
    resolve_embedding_model,
    store_embedding_records,
)
from app.services.knowledge_asset_service import KnowledgeAssetService, register_document_asset as _register_doc_asset, register_audio_asset as _register_audio_asset, register_video_asset as _register_video_asset
from app.db.models import EMBEDDING_VERSION
from datetime import datetime

logger = logging.getLogger(__name__)

# Cache for pre-chunked audio/video segments from process_audio_for_rag().
# Keyed by file path → {"chunks": list[dict], "source_type": str}
# Set by extract_text(), consumed by run_embedding_pipeline().
_audio_chunks_cache: dict[str, dict] = {}
# Cache for audio/video metadata from process_audio_for_rag().
# Keyed by file path -> {"duration_sec": float, "word_count": int, "model_used": str, ...}
# Set by extract_text(), consumed by ingest_document() after embedding pipeline.
_audio_meta_cache: dict[str, dict] = {}
# Cache for video-specific metadata from video_analysis_service.
# Keyed by file path -> {"width": int, "height": int, "fps": float, "codec": str, ...}
# Set by extract_text() for video files.
_video_meta_cache: dict[str, dict] = {}

def _now_iso() -> str:
    return time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime())

_STATE_MAP = {
    "uploaded": "UPLOADED",
    "ingesting": "PROCESSING",
    "embedded": "PROCESSING",
    "summarized": "PROCESSING",
    "completed": "COMPLETED",
    "failed": "FAILED",
}

_SAVE_CHECKPOINT = True


def _make_checkpoint(doc_id, step: str, percent: int, detail: str = "") -> str:
    return json.dumps({
        "doc_id": str(doc_id),
        "step": step,
        "percent": percent,
        "detail": detail,
        "ts": _now_iso(),
    })


async def _set_doc_state(
    db: AsyncSession,
    doc: Document,
    *,
    status: str,
    step: str,
    percent: int,
    message: str = "",
    error_message: str = "",
) -> None:
    doc.status = status
    doc.ingest_step = step
    doc.ingest_percent = percent
    doc.ingest_message = message
    doc.error_message = error_message
    doc.processing_step = step
    doc.processing_state = _STATE_MAP.get(status, "PROCESSING")
    doc.updated_at = datetime.now(timezone.utc)
    if status == "uploaded":
        doc.ingestion_started = False
        doc.ingestion_completed = False
        doc.ingestion_failed = False
        doc.analysis_ready = False
    elif status in ("ingesting", "embedded", "summarized"):
        doc.ingestion_started = True
        doc.ingestion_failed = False
    elif status == "completed":
        doc.ingestion_started = True
        doc.ingestion_completed = True
        doc.ingestion_failed = False
        doc.analysis_ready = True
    elif status == "failed":
        doc.ingestion_started = True
        doc.ingestion_failed = True

    # ── Save checkpoint for resume ──
    if _SAVE_CHECKPOINT:
        doc.checkpoint = _make_checkpoint(doc.id, step, percent, message)

    db.add(doc)

    try:
        await db.commit()
        # Re-apply RLS context after commit — the underlying connection
        # may have changed, which would lose the SET SESSION setting.
        if doc.owner_id:
            await db.execute(
                text("SELECT set_config('app.current_user_id', :val, false)"),
                {"val": str(doc.owner_id)},
            )
    except Exception:
        try:
            await db.rollback()
        except Exception:
            pass
        raise

async def get_file_hash(file_path: str) -> str:
    sha256_hash = hashlib.sha256()
    with open(file_path, "rb") as f:
        for byte_block in iter(lambda: f.read(4096), b""):
            sha256_hash.update(byte_block)
    return sha256_hash.hexdigest()

async def extract_text(file_path: str, mime_type: str) -> str:
    text = ""
    file_ext = Path(file_path).suffix.lower()
    logger.info(f"[EXTRACT] Starting text extraction for {file_path}")
    logger.info(f"[EXTRACT] mime_type={mime_type}, file_ext={file_ext}")
    
    if file_ext == ".pdf":
        try:
            # ── Tier 1: Try pdfplumber (structured text + table extraction) ──
            # pdfplumber preserves table structure (rows → cells) that PyPDF2
            # loses by flattening everything into running text. This is critical
            # for documents with requirements tables, spec matrices, etc.
            try:
                import pdfplumber
                page_texts: list[str] = []
                table_count = 0
                with pdfplumber.open(file_path) as pdf:
                    logger.info(f"[EXTRACT] PDF has {len(pdf.pages)} pages (via pdfplumber)")
                    for page in pdf.pages:
                        # Extract regular text first
                        page_text = page.extract_text() or ""
                        parts = [page_text] if page_text.strip() else []
                        # Extract tables with structure preserved
                        tables = page.extract_tables()
                        for table in tables:
                            if not table or len(table) < 2:
                                continue
                            table_count += 1
                            rows: list[str] = []
                            for row in table:
                                cells = [str(cell or "").strip() for cell in row]
                                rows.append(" | ".join(cells))
                            parts.append("\n" + "\n".join(rows) + "\n")
                        page_texts.append("\n".join(parts))
                text = "\n".join(page_texts)
                logger.info(
                    "[EXTRACT] pdfplumber extraction: %d chars, %d tables across %d pages",
                    len(text), table_count, len(pdf.pages),
                )
            except ImportError:
                # pdfplumber not installed — fall through to PyPDF2
                logger.info("[EXTRACT] pdfplumber not available; falling back to PyPDF2")
                text = ""
                reader = PdfReader(file_path)
                for i, page in enumerate(reader.pages):
                    page_text = page.extract_text() or ""
                    text += page_text
                logger.info("[EXTRACT] PyPDF2 fallback: %d chars", len(text))
            
            extracted_len = len(text)
            stripped_len = len(text.strip())
            
            # Fallback to OCR if text is poor/empty (for scanned/image-based PDFs)
            if stripped_len < 50:
                logger.warning(f"[EXTRACT] PDF text extraction poor ({stripped_len} chars), attempting OCR fallback")
                if _check_tesseract():
                    logger.info(f"[EXTRACT] OCR fallback enabled, processing PDF with pypdfium2")
                    try:
                        import pypdfium2 as pdfium
                    except Exception as e:
                        raise RuntimeError("PDF OCR requires pypdfium2") from e
                    pdf = pdfium.PdfDocument(file_path)
                    ocr_parts: list[str] = []
                    for i in range(len(pdf)):
                        page = pdf.get_page(i)
                        pil_image = page.render(scale=2).to_pil()
                        try:
                            ocr_text = pytesseract.image_to_string(pil_image)
                            ocr_parts.append(ocr_text)
                            logger.info(f"[EXTRACT] OCR page {i+1}: {len(ocr_text)} chars")
                        except Exception as ocr_err:
                            logger.warning(f"[EXTRACT] OCR failed for page {i+1}: {ocr_err}")
                        page.close()
                    pdf.close()
                    text = "\n".join(ocr_parts)
                    logger.info(f"[EXTRACT] OCR complete: {len(text)} chars total")
                else:
                    logger.error(f"[EXTRACT] OCR not available - tesseract not found in PATH")
        except Exception as e:
            logger.error(f"[EXTRACT] Error extracting text from PDF {file_path}: {e}", exc_info=True)
            
    elif file_ext == ".docx":
        try:
            doc = docx.Document(file_path)
            # Extract paragraphs (old behavior) — always works
            para_texts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
            # Extract tables — doc.paragraphs does NOT include table content!
            # python-docx stores tables separately; we must iterate doc.tables explicitly.
            # Note: We cannot use id() to match doc.element.body children to doc.paragraphs/tables
            # because lxml creates different proxy objects for the same element at different access
            # times, making id() comparison unreliable. Instead, we extract in two passes:
            # first all paragraphs, then all tables, concatenated in order.
            table_texts: list[str] = []
            for table in doc.tables:
                rows: list[str] = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(" | ".join(cells))
                table_texts.append("\n".join(rows))
            # Combine paragraphs + tables with clear separation
            all_text = para_texts + table_texts
            text = "\n".join(all_text)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {e}")
            raise
            
    elif file_ext in [".txt", ".md"]:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
        except Exception as e:
            logger.error(f"Error reading TXT {file_path}: {e}")
            
    elif file_ext in [".png", ".jpg", ".jpeg"]:
        if _check_tesseract():
            try:
                text = pytesseract.image_to_string(Image.open(file_path))
            except Exception as e:
                logger.error(f"Error performing OCR on image {file_path}: {e}")
        else:
            logger.info(f"Skipping OCR for {file_path} – tesseract not available")

    elif file_ext in {".wav", ".mp3", ".m4a", ".ogg", ".flac", ".aac", ".wma"}:
        try:
            from app.services.transcription_service import process_audio_for_rag
            result = await process_audio_for_rag(file_path, mime_type=mime_type)
            rag_chunks: list[dict] = result.get("rag_chunks", [])
            full_text: str = result.get("full_text", "")
            stype: str = result.get("source_type", "audio")
            if rag_chunks:
                _audio_chunks_cache[file_path] = {"chunks": rag_chunks, "source_type": stype}
                # Also cache metadata for post-ingestion storage
                _audio_meta_cache[file_path] = {
                    "duration_sec": result.get("duration_sec"),
                    "word_count": result.get("word_count"),
                    "model_used": result.get("model_used"),
                    "language": result.get("language"),
                    "source_type": stype,
                }
                text = " ".join(c["text"] for c in rag_chunks)
                logger.info(f"Audio pipeline: {len(rag_chunks)} natural chunks, source_type={stype}, {len(text)} chars from {file_path}")
            elif full_text:
                text = full_text
                logger.info(f"Audio transcription extracted {len(text)} chars from {file_path} (no chunks)")
            else:
                from app.services.transcription_service import diagnose_transcription_failure, _set_last_error
                diag = diagnose_transcription_failure()
                _set_last_error(diag.get("message", "Audio transcription returned no content"))
                logger.warning("Audio transcription returned no content for %s — %s", file_path, diag.get("message"))
        except Exception as e:
            logger.error(f"Error transcribing audio {file_path}: {e}")

    elif file_ext in {".mp4", ".avi", ".mov", ".mkv", ".wmv", ".flv"}:
        try:
            from app.services.transcription_service import process_audio_for_rag
            result = await process_audio_for_rag(file_path, mime_type=mime_type)
            rag_chunks: list[dict] = result.get("rag_chunks", [])
            full_text: str = result.get("full_text", "")
            stype: str = result.get("source_type", "video")
            if rag_chunks:
                _audio_chunks_cache[file_path] = {"chunks": rag_chunks, "source_type": stype}
                text = " ".join(c["text"] for c in rag_chunks)
                logger.info(f"Video pipeline: {len(rag_chunks)} natural chunks, source_type={stype}, {len(text)} chars from {file_path}")
            elif full_text:
                text = full_text
                logger.info(f"Video transcription extracted {len(text)} chars from {file_path} (no chunks)")
            else:
                from app.services.transcription_service import diagnose_transcription_failure, _set_last_error
                diag = diagnose_transcription_failure()
                _set_last_error(diag.get("message", "Video transcription returned no content"))
                logger.warning("Video transcription returned no content for %s — %s", file_path, diag.get("message"))

            # Also run lightweight video analysis for metadata (duration, dimensions, codec)
            try:
                from app.services.video_analysis_service import analyze_video_light
                video_info = await analyze_video_light(file_path, mime_type=mime_type)
                if video_info:
                    _video_meta_cache[file_path] = video_info
                    logger.info(
                        "Video metadata: %.0fs, %dx%d, %.1ffps, %s",
                        video_info.get("duration_sec", 0),
                        video_info.get("width", 0),
                        video_info.get("height", 0),
                        video_info.get("fps", 0),
                        video_info.get("codec", "unknown"),
                    )
            except Exception as ve:
                logger.warning("Video metadata extraction skipped: %s", ve)
        except Exception as e:
            logger.error(f"Error transcribing video {file_path}: {e}")

    return text

def chunk_text(text: str, chunk_size: int, chunk_overlap: int) -> List[str]:
    chunks = []
    # Safeguard against infinite loops and bad configs
    if chunk_size <= 0:
        chunk_size = 1000
    if chunk_overlap >= chunk_size:
        chunk_overlap = chunk_size // 10

    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        if end >= len(text):
            break
        start += (chunk_size - chunk_overlap)
    return chunks

async def analyze_document(text: str, doc_id: int, db: AsyncSession, filename: str = "", detected_type: str = ""):
    # Cap context so even small models don't OOM. 6 000 chars ≈ 1 500 tokens.
    analysis_text = text[:6000].strip()

    # ── Resolve model via centralized model resolver ──────────────────────
    # This consults UI-configured task mappings as the single source of truth.
    resolved = await resolve_model(db, requested_model=None, task_type="summary")
    model_name = resolved["model_id"]

    # ── Single combined prompt ──────────────────────────────────────────────
    # One round-trip to Ollama or Gemini produces everything we need.
    combined_prompt = (
        "You are a professional document analyst for ZETDC (Zimbabwe Electricity Transmission and Distribution Company). "
        "Analyze the document excerpt below and return ONLY a valid JSON object "
        "(no markdown, no commentary, no asterisks) with exactly these keys:\n"
        "  executive_summary  – string: a professional narrative summary in flowing prose (max 10 sentences). "
        "Do NOT use asterisks, bold markdown, numbered lists, or bullet points. Write as well-structured paragraphs.\n"
        "  detailed_summary   – array of strings: each string must be a complete narrative sentence (not a bullet point). "
        "Present key findings as professional prose, not as a list of fragments.\n"
        "  sentiment          – one of: Positive, Neutral, Negative, Urgent\n"
        "  topics             – array of short strings (max 8)\n"
        "  entities           – array of strings (people, orgs, systems)\n"
        "  action_items       – array of strings: each as a complete professional sentence\n"
        "  decisions          – array of strings: each as a complete professional sentence\n"
        "  prompts            – array of 5 specific, actionable questions a user could ask "
        "about this document (include one Mermaid diagram prompt and one action-items prompt)\n\n"
        "SUMMARY WRITING RULES:\n"
        "- NEVER use asterisks (**bold**) or markdown formatting for emphasis\n"
        "- NEVER use numbered or bulleted lists in executive_summary; write in narrative paragraphs\n"
        "- Each detailed_summary entry must be a complete, well-formed sentence — not a fragment or bullet point\n"
        "- Maintain a formal, professional tone suitable for ZETDC leadership and staff\n"
        "- Begin executive_summary with a clear statement of the document's purpose and scope\n"
        "- Organise content logically: context first, then key findings, then implications\n"
        "- Use precise ZETDC terminology (transmission, distribution, substations, feeders, SCADA, HSE, ZERA compliance)\n\n"
        "Document excerpt:\n" + analysis_text
    )

    structured: dict = {}
    
    # ── Try Gemini API first (superior analysis quality) ───────────────────
    if settings.gemini_api_key:
        try:
            from app.services.gemini_service import generate as gemini_generate
            logger.info(f"Analyzing document {doc_id} with Gemini API")
            raw = await gemini_generate(
                combined_prompt,
                system="You are a precise document analyst. Output only valid JSON."
            )
            json_start = raw.find("{")
            json_end = raw.rfind("}") + 1
            if json_start >= 0 and json_end > json_start:
                structured = json.loads(raw[json_start:json_end])
                logger.info(f"Document {doc_id}: Gemini analysis successful")
        except Exception as e:
            logger.warning(f"Gemini analysis failed for doc {doc_id} ({e}); trying DeepSeek API")
    
    # ── Try DeepSeek API if Gemini failed ──────────────────────────────────
    if not structured and settings.deepseek_api_key:
        try:
            from app.services.deepseek_service import generate as deepseek_generate
            logger.info(f"Analyzing document {doc_id} with DeepSeek API")
            raw = await deepseek_generate(
                combined_prompt,
                system="You are a precise document analyst. Output only valid JSON."
            )
            json_start = raw.find("{")
            json_end = raw.rfind("}") + 1
            if json_start >= 0 and json_end > json_start:
                structured = json.loads(raw[json_start:json_end])
                logger.info(f"Document {doc_id}: DeepSeek analysis successful")
        except Exception as e:
            logger.warning(f"DeepSeek analysis failed for doc {doc_id} ({e}); falling back to Ollama")
    
    # ── Fallback to Ollama local model ──────────────────────────────────────
    if not structured:
        try:
            logger.info(f"Analyzing document {doc_id} with Ollama model: {model_name}")
            raw = await ollama.generate(
                model_name,
                combined_prompt,
                system="You are a precise document analyst. Output only valid JSON.",
                options={"num_ctx": 4096, "temperature": 0},
            )
            json_start = raw.find("{")
            json_end = raw.rfind("}") + 1
            if json_start >= 0 and json_end > json_start:
                structured = json.loads(raw[json_start:json_end])
                logger.info(f"Document {doc_id}: Ollama analysis successful")
        except Exception as e:
            logger.warning(f"Ollama analysis failed for doc {doc_id} ({e}); using extractive fallback")

    # ── Extractive fallback ─────────────────────────────────────────────────
    def _extract_sentences(t: str, n: int) -> list[str]:
        sents = [s.strip() for s in re.split(r"(?<=[.!?])\s+", t.replace("\n", " ")) if s.strip()]
        return sents[:n]

    def _extract_entities_simple(t: str) -> list[str]:
        words = re.findall(r"\b[A-Z][a-zA-Z]{2,}(?:\s[A-Z][a-zA-Z]{2,})*\b", t)
        seen: dict[str, int] = {}
        for w in words:
            seen[w] = seen.get(w, 0) + 1
        return [w for w, _ in sorted(seen.items(), key=lambda x: -x[1])][:10]

    def _infer_sentiment_simple(t: str) -> str:
        tl = t.lower()
        if any(w in tl for w in ("urgent", "critical", "immediate", "emergency", "risk", "fail")):
            return "Urgent"
        if any(w in tl for w in ("success", "achiev", "complet", "positive", "benefit", "improve")):
            return "Positive"
        if any(w in tl for w in ("problem", "issue", "error", "negativ", "concern", "loss")):
            return "Negative"
        return "Neutral"

    default_prompts = [
        "Summarize this document in 10 bullets.",
        "List all action items and decisions mentioned in this document.",
        "Generate a Mermaid process flow diagram from the key steps in this document.",
        "Extract all key entities: people, departments, locations, systems, and dates.",
        "What are the key risks, mitigations, and compliance implications?",
    ]

    exec_summary = str(structured.get("executive_summary") or "").strip()
    if not exec_summary:
        sentences = _extract_sentences(analysis_text, 10)
        exec_summary = " ".join(sentences) if sentences else "No content available."

    def _ensure_list(val, default: list) -> list:
        if isinstance(val, list) and val:
            return [str(x).strip() for x in val if str(x).strip()]
        return default

    detailed_summary = _ensure_list(structured.get("detailed_summary"), _extract_sentences(analysis_text, 5))
    entities = _ensure_list(structured.get("entities"), _extract_entities_simple(analysis_text))
    topics = _ensure_list(structured.get("topics"), [])
    action_items = _ensure_list(structured.get("action_items"), [])
    decisions = _ensure_list(structured.get("decisions"), [])
    suggested_prompts = _ensure_list(structured.get("prompts"), default_prompts)
    if len(suggested_prompts) < 3:
        suggested_prompts = (suggested_prompts + default_prompts)[:5]

    sentiment = str(structured.get("sentiment") or "").strip().title()
    if sentiment not in {"Positive", "Neutral", "Negative", "Urgent"}:
        sentiment = _infer_sentiment_simple(analysis_text)

    # ── Generate enterprise summary (document-type-aware) ────────────────
    enterprise_summary: dict | None = None
    try:
        from app.services.document_summarizer import generate_enterprise_summary
        # Determine doc_type from explicit detected_type or infer from analysis
        detected_doc_type = detected_type or ""
        if not detected_doc_type:
            if is_meeting_like(topics, entities):
                detected_doc_type = "meeting"
            elif is_policy_like(topics, entities):
                detected_doc_type = "policy"

        # Pass the full text — the summarizer handles its own input cap (12000 chars)
        enterprise_summary = await generate_enterprise_summary(
            db,
            text,  # full extracted text; summarizer truncates internally
            filename=filename or "",
            detected_type=detected_doc_type,
        )
        logger.info(
            "[SUMMARY] Enterprise summary generated for doc %s (type=%s)",
            doc_id, enterprise_summary.get("doc_type", "unknown"),
        )
    except Exception as e:
        logger.warning("[SUMMARY] Enterprise summary generation failed for doc %s: %s", doc_id, e)
        enterprise_summary = None

    # ── Persist to DB ───────────────────────────────────────────────────────
    analysis = DocAnalysis(
        document_id=doc_id,
        executive_summary=exec_summary,
        detailed_summary="\n".join(detailed_summary),
        sentiment=sentiment,
        entities_json=json.dumps(entities),
        topics_json=json.dumps(topics),
        action_items_json=json.dumps(action_items),
        decisions_json=json.dumps(decisions),
        summary_json=json.dumps(enterprise_summary) if enterprise_summary else None,
        doc_type=enterprise_summary.get("doc_type", "") if enterprise_summary else None,
    )
    db.add(analysis)

    for p_text in suggested_prompts[:5]:
        db.add(SuggestedPrompt(document_id=doc_id, prompt_text=p_text))

    await db.commit()


def is_meeting_like(topics: list[str], entities: list[str]) -> bool:
    """Heuristic check if document appears to be a meeting transcript."""
    lower_topics = " ".join(t.lower() for t in topics)
    lower_entities = " ".join(e.lower() for e in entities)
    combined = lower_topics + " " + lower_entities
    keywords = ["meeting", "transcript", "discussion", "participant", "agenda",
                "minutes", "decision", "action item", "speaker"]
    return any(kw in combined for kw in keywords)


def is_policy_like(topics: list[str], entities: list[str]) -> bool:
    """Heuristic check if document appears to be a policy."""
    lower_topics = " ".join(t.lower() for t in topics)
    lower_entities = " ".join(e.lower() for e in entities)
    combined = lower_topics + " " + lower_entities
    keywords = ["policy", "compliance", "regulation", "governance", "procedure",
                "responsibility", "obligation", "control", "audit"]
    score = sum(2 for kw in keywords if kw in combined)
    return score >= 4

async def ingest_document(doc_id: uuid.UUID, db: AsyncSession):
    # ── DIAG: Check RLS context before SELECT Document ──
    rls_check = await db.execute(
        text("SELECT current_setting('app.current_user_id', true)")
    )
    rls_val = rls_check.scalar()
    logger.info(f"[DIAG:INGEST] ingest_document called for doc_id={doc_id} RLS context={rls_val!r}")

    # Get document details
    result = await db.execute(select(Document).where(Document.id == doc_id))
    doc = result.scalar_one_or_none()
    if not doc:
        logger.error(f"[DIAG:INGEST] Document {doc_id} NOT FOUND by SELECT — RLS context was {rls_val!r}")
        return

    # Log document metadata at start
    import os
    file_size = 0
    try:
        file_size = os.path.getsize(doc.path) if doc.path else 0
    except Exception:
        pass
    
    logger.info(f"[INGEST] {'='*60}")
    logger.info(f"[INGEST] Starting ingestion for document_id={doc_id}")
    logger.info(f"[INGEST] filename={doc.filename}")
    logger.info(f"[INGEST] file_size={file_size} bytes")
    logger.info(f"[INGEST] mime_type={doc.mime_type}")
    logger.info(f"[INGEST] project_id={doc.project_id}")
    logger.info(f"[INGEST] path={doc.path}")

    try:
        await _set_doc_state(db, doc, status="ingesting", step="extract", percent=0, message="Extracting text")
        
        # TEXT EXTRACTION
        logger.info(f"[INGEST] Extracting text from {doc.path}")
        extracted_text = await extract_text(doc.path, doc.mime_type)
        extracted_length = len(extracted_text) if extracted_text else 0
        stripped_length = len(extracted_text.strip()) if extracted_text else 0
        
        logger.info(f"[INGEST] Extracted {extracted_length} characters from document {doc_id}")
        logger.info(f"[INGEST] After stripping whitespace: {stripped_length} characters")
        
        if not extracted_text or stripped_length < 10:
            logger.error(f"[INGEST] Text extraction failed or produced insufficient text for document {doc_id}")
            logger.error(f"[INGEST] Extracted text sample: {repr(extracted_text[:200] if extracted_text else '(empty)')}")
            await _set_doc_state(
                db,
                doc,
                status="failed",
                step="extract",
                percent=0,
                message="No extractable text",
                error_message=f"No text could be extracted from the document (extracted: {extracted_length} chars, usable: {stripped_length} chars)",
            )
            return

        await _set_doc_state(db, doc, status="ingesting", step="chunk", percent=20, message="Chunking text")
        
        # Log first 500 chars of extracted text for debugging
        logger.info(f"[INGEST] Text sample for document {doc_id}: {repr(extracted_text[:500])}")

        # Initial chunking (e.g., ~1000 chars = ~250 tokens)
        init_chunk_size = settings.chunk_size
        init_overlap = settings.chunk_overlap

        # ── Resolve embedding model for governance tracking ────────────────
        embedding_model_info = await resolve_embedding_model(db)
        embed_provider = embedding_model_info["provider_name"] if embedding_model_info else ""
        embed_model_id = embedding_model_info["model_id"] if embedding_model_info else ""

        async def run_embedding_pipeline(c_size: int, c_overlap: int) -> tuple[list[str], list[list[float]], list[str], list[dict]]:
            # Check for pre-chunked audio/video segments from extract_text()
            cached = _audio_chunks_cache.pop(doc.path, None)
            pre_chunked = cached["chunks"] if cached else None
            pre_source_type = cached["source_type"] if cached else None
            if pre_chunked:
                chunks = [c["text"] for c in pre_chunked]
                chunks = [c.strip() for c in chunks if c.strip()]
                logger.info(f"[INGEST] Using {len(chunks)} pre-chunked audio/video segments for document {doc_id}")
                if not chunks:
                    raise ValueError("Audio chunks produced no text segments")
            else:
                logger.info(f"[INGEST] Chunking text with chunk_size={c_size}, overlap={c_overlap}")
                chunks = chunk_text(extracted_text, c_size, c_overlap)
                chunks = [c.strip() for c in (chunks or []) if isinstance(c, str) and c.strip()]
                logger.info(f"[INGEST] Created {len(chunks)} chunks for document {doc_id}")
                if not chunks:
                    raise ValueError("Chunking produced no chunks")

            sem = asyncio.Semaphore(2)
            async def embed_one(i: int, chunk_text_content: str):
                async with sem:
                    try:
                        vec = await generate_embedding(db, chunk_text_content, embed_model_id)
                        return i, vec
                    except Exception as e:
                        logger.error(f"Embedding failed at chunk {i}: {e}")
                        raise

            results = await asyncio.gather(*(embed_one(i, c) for i, c in enumerate(chunks)))
            results.sort(key=lambda x: x[0])
            
            emb_out, metadatas, ids, docs = [], [], [], []
            for i, vec in results:
                if i < 0 or i >= len(chunks):
                    continue
                content = chunks[i].strip()
                if not content:
                    continue
                chroma_id = f"chroma_{doc_id}_{i}_{c_size}"
                ids.append(chroma_id)
                emb_out.append(vec)
                docs.append(content)
                # Determine source_type from file extension
                # For pre-chunked audio/video, use the source_type from the pipeline
                f_ext = Path(doc.filename).suffix.lower() if doc.filename else ""
                if pre_source_type:
                    stype = pre_source_type
                elif f_ext in [".pdf", ".docx", ".txt", ".md"]:
                    stype = "document"
                elif f_ext in [".png", ".jpg", ".jpeg"]:
                    stype = "image"
                elif f_ext in [".wav", ".mp3", ".m4a", ".ogg", ".flac", ".aac", ".wma"]:
                    stype = "audio"
                elif f_ext in [".mp4", ".avi", ".mov", ".mkv", ".wmv", ".flv"]:
                    stype = "video"
                else:
                    stype = "unknown"
                # Build metadata – include start_sec/end_sec for audio/video segments
                meta: dict = {
                    "filename": doc.filename,
                    "chunk_index": i,
                    "document_id": str(doc_id),
                    "source_type": stype,
                }
                if pre_chunked and i < len(pre_chunked):
                    seg = pre_chunked[i]
                    if "start_sec" in seg:
                        meta["start_sec"] = seg["start_sec"]
                    if "end_sec" in seg:
                        meta["end_sec"] = seg["end_sec"]
                    if "speaker" in seg and seg["speaker"]:
                        meta["speaker"] = seg["speaker"]
                metadatas.append(meta)
            return chunks, emb_out, ids, docs, metadatas

        await _set_doc_state(db, doc, status="embedded", step="embed", percent=45, message="Generating embeddings")
        
        try:
            _, embeddings, ids, documents, metadatas = await run_embedding_pipeline(init_chunk_size, init_overlap)
        except Exception as embed_err:
            err_str = str(embed_err).lower()
            if "token out of range" in err_str or "context length" in err_str or "size limit" in err_str:
                logger.warning(f"Token limit exceeded for doc {doc_id}. Retrying with smaller chunks (50% size).")
                await _set_doc_state(db, doc, status="embedded", step="embed", percent=45, message="Recovering from token limit; retrying with smaller chunks")
                fallback_size = max(200, init_chunk_size // 2)
                fallback_overlap = max(20, init_overlap // 2)
                _, embeddings, ids, documents, metadatas = await run_embedding_pipeline(fallback_size, fallback_overlap)
            else:
                raise embed_err

        # CRITICAL: Fail ingestion if no chunks were produced
        chunk_count = len(ids) if ids else 0
        if chunk_count == 0:
            logger.error(f"[INGEST] ✗ Document {doc_id} produced ZERO chunks. Ingestion FAILED.")
            logger.error(f"[INGEST] Text length was {len(extracted_text)} chars, but chunking produced no valid chunks.")
            await _set_doc_state(
                db,
                doc,
                status="failed",
                step="chunk",
                percent=20,
                message="Chunk generation failed",
                error_message=f"Document produced no chunks. The document may be empty, corrupted, or contain no extractable text. Text extracted: {len(extracted_text)} characters.",
            )
            return
        
        logger.info(f"[INGEST] Document {doc_id} produced {chunk_count} chunks successfully")

        if not (len(ids) == len(embeddings) == len(documents) == len(metadatas)):
            raise RuntimeError(
                f"Vector upsert input mismatch: ids={len(ids)} docs={len(documents)} emb={len(embeddings)} meta={len(metadatas)}"
            )

        # ── DUAL-WRITE: ChromaDB + pgvector ────────────────────────────────
        # Always write to ChromaDB (primary) for backward compatibility.
        # When USE_PGVECTOR is enabled, also write to pgvector (document_chunks).

        batch_size = 96
        failures: list[str] = []

        for start in range(0, len(ids), batch_size):
            end = min(len(ids), start + batch_size)
            bid = ids[start:end]
            bdoc = documents[start:end]
            bemb = embeddings[start:end]
            bmeta = metadatas[start:end]
            if not (len(bid) == len(bdoc) == len(bemb) == len(bmeta)) or len(bid) == 0:
                failures.append(f"batch_len_mismatch({start}-{end})")
                continue
            # ── ChromaDB write (always) ────────────────────────────────────
            print(f"CHROMA_UPSERT: calling with {len(bid)} ids, {len(bdoc)} docs, {len(bemb)} emb, {len(bmeta)} meta for project_{doc.project_id}", flush=True)
            try:
                chroma.upsert(str(doc.project_id), bid, bdoc, bemb, bmeta)
                print(f"CHROMA_UPSERT: SUCCESS for project_{doc.project_id}", flush=True)
            except Exception as e:
                print(f"CHROMA_UPSERT: FAILED with {e}", flush=True)
                failures.append(str(e))
                break

        if failures:
            raise RuntimeError("Embedding upsert failed: " + "; ".join(failures[:3]))

        # ── pgvector dual-write (when feature-flagged) ─────────────────────
        if app_settings.use_pgvector:
            pg_success = False
            pg_error = ""
            try:
                # Collect texts and metadata for pgvector insertion
                pg_texts = []
                pg_metas = []
                for i in range(len(ids)):
                    pg_texts.append(documents[i])
                    pg_metas.append(metadatas[i] if i < len(metadatas) else {})

                inserted = await pgvector_insert_chunks(
                    db,
                    document_id=doc.id,
                    texts=pg_texts,
                    embeddings=embeddings,
                    metadatas=pg_metas,
                )
                pg_success = inserted > 0
                logger.info(
                    "[PGVECTOR] Dual-write: inserted %d chunks for document %s",
                    inserted, doc.id,
                )
            except Exception as e:
                pg_error = str(e)
                logger.error("[PGVECTOR] Dual-write FAILED for document %s: %s", doc.id, e, exc_info=True)

            if not pg_success and pg_error:
                # pgvector write failed but ChromaDB succeeded — log and continue
                logger.warning(
                    "[PGVECTOR] Dual-write failed but ChromaDB succeeded for doc %s: %s",
                    doc.id, pg_error,
                )

        # ── CHROMA VERIFICATION: check count immediately after upsert ──────
        try:
            c_before = chroma.get_collection(str(doc.project_id)).count()
            print(f"CHROMA_VERIFY: project_{doc.project_id} count={c_before}", flush=True)
        except Exception as ce:
            print(f"CHROMA_VERIFY: error getting count: {ce}", flush=True)

        # ── Compute dimensions & store governance records ──────────────────
        embed_dimensions = len(embeddings[0]) if embeddings else 0
        logger.info(f"[INGEST] Storing {len(ids)} embedding records for document {doc_id}")
        await store_embedding_records(
            db, doc_id, doc.project_id, ids, documents, metadatas,
            provider=embed_provider, model=embed_model_id, dimensions=embed_dimensions,
        )
        logger.info(f"[INGEST] Successfully stored embedding records for document {doc_id}")

        # ── Update Document with embedding governance metadata ─────────────
        doc.embedding_provider = embed_provider or None
        doc.embedding_model = embed_model_id or None
        doc.embedding_version = EMBEDDING_VERSION
        doc.embedded_at = datetime.utcnow()
        await db.commit()
        logger.info(f"[INGEST] Document {doc_id} embedding metadata updated: provider={embed_provider}, model={embed_model_id}")

        # ── Knowledge Asset Registry: Register document and chunks ────────
        try:
            registry = KnowledgeAssetService(db)
            await registry.register_document_asset(doc)
            # Re-fetch chunks to get their IDs for registration
            chunk_result = await db.execute(
                select(Chunk).where(Chunk.document_id == doc.id).order_by(Chunk.chunk_index)
            )
            db_chunks = list(chunk_result.scalars().all())
            if db_chunks:
                await registry.register_chunks(doc.id, db_chunks)
            # Register embeddings
            emb_result = await db.execute(
                select(Embedding).where(
                    Embedding.vector_ref.in_([f"chroma_{doc_id}_{i}_{settings.chunk_size}" for i in range(len(ids))])
                )
            )
            db_embeddings = list(emb_result.scalars().all())
            logger.info("[ASSET] Registered document %s and %d chunks via KnowledgeAssetService",
                        doc.id, len(db_chunks))
        except Exception as asset_err:
            logger.warning("[ASSET] Failed to register assets for document %s: %s", doc.id, asset_err)

        # ── Store audio/video metadata if applicable ───────────────────────
        _audio_meta_for_asset = _audio_meta_cache.pop(doc.path, None)
        _video_meta_for_asset = _video_meta_cache.pop(doc.path, None)

        if _audio_meta_for_asset:
            try:
                from app.db.models import AudioMetadata
                meta_record = AudioMetadata(
                    document_id=doc.id,
                    duration_sec=_audio_meta_for_asset.get("duration_sec"),
                    word_count=_audio_meta_for_asset.get("word_count"),
                    transcription_model=_audio_meta_for_asset.get("model_used"),
                    language=_audio_meta_for_asset.get("language", "en"),
                    source_type=_audio_meta_for_asset.get("source_type", "audio"),
                    processing_time_ms=None,
                )
                db.add(meta_record)
                await db.commit()
                logger.info("[AUDIO_META] Stored audio metadata for document %s: duration=%s, model=%s",
                            doc.id, _audio_meta_for_asset.get("duration_sec"), _audio_meta_for_asset.get("model_used"))
            except Exception as meta_err:
                logger.warning("[AUDIO_META] Failed to store audio metadata for doc %s: %s", doc.id, meta_err)

        if _video_meta_for_asset:
            logger.info(
                "[VIDEO_META] Video metadata for doc %s: %.0fs, %dx%d, %.1ffps, %s",
                doc.id,
                _video_meta_for_asset.get("duration_sec", 0),
                _video_meta_for_asset.get("width", 0),
                _video_meta_for_asset.get("height", 0),
                _video_meta_for_asset.get("fps", 0),
                _video_meta_for_asset.get("codec", "unknown"),
            )

        await _set_doc_state(db, doc, status="summarized", step="summarize", percent=80, message="Generating summaries")
        await analyze_document(extracted_text, doc_id, db, filename=doc.filename or "", detected_type=doc.detected_type or "")

        # ── Knowledge Asset Registry: Register analysis after creation ────
        try:
            analysis_result = await db.execute(
                select(DocAnalysis).where(DocAnalysis.document_id == doc.id)
            )
            db_analysis = analysis_result.scalar_one_or_none()
            if db_analysis:
                # Analysis registered via register_document_asset path
                logger.info("[ASSET] Analysis %s registered for document %s", db_analysis.id, doc.id)
        except Exception as asset_err:
            logger.warning("[ASSET] Failed to register analysis for document %s: %s", doc.id, asset_err)

        # ── Register as Knowledge Asset (upgraded service) ───────────────
        try:
            analysis_result = await db.execute(
                select(DocAnalysis).where(DocAnalysis.document_id == doc.id)
            )
            db_analysis = analysis_result.scalar_one_or_none()
            # Use already-popped metadata variables (not re-reading cache)
            file_ext = Path(doc.filename).suffix.lower() if doc.filename else ""
            if file_ext in {".mp4", ".avi", ".mov", ".mkv", ".wmv", ".flv"} and _video_meta_for_asset:
                # Register as video asset with full video metadata + transcript
                transcript = _video_meta_for_asset.get("transcript", "")
                await _register_video_asset(db, doc, _video_meta_for_asset, transcript=transcript)
            elif _audio_meta_for_asset:
                await _register_audio_asset(db, doc, _audio_meta_for_asset)
            await _register_doc_asset(db, doc, db_analysis)
            logger.info("[ASSET] Upgraded asset registration for document %s", doc.id)
        except Exception as asset_err:
            logger.warning("[ASSET] KnowledgeAsset registration failed for doc %s: %s", doc.id, asset_err)

        await _set_doc_state(db, doc, status="completed", step="done", percent=100, message="Completed")
        logger.info(f"[INGEST] ✓ Document {doc_id} ingestion COMPLETED successfully")
        logger.info(f"[INGEST] {'='*60}")
        # Auto-trigger transfer learning on local models
        _schedule_auto_training(doc.project_id or 0)
    except Exception as e:
        await _set_doc_state(
            db,
            doc,
            status="failed",
            step="failed",
            percent=0,
            message="Failed",
            error_message=str(e),
        )
