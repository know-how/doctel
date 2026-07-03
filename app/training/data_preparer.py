"""
data_preparer.py – Scans training_room/inbox/ for supported documents and converts
them to JSONL prompt-completion pairs ready for LoRA fine-tuning.

Supported input formats: .pdf, .docx, .txt, .log, .csv, .md
Output: training_room/batches/batch_<timestamp>.jsonl

Each JSONL line is a dict:
  {"prompt": "<text excerpt>", "completion": "<continuation or summary>"}
"""
import json
import logging
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Generator

logger = logging.getLogger(__name__)

# ── text extraction helpers ────────────────────────────────────────────────────

def _extract_txt(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8", errors="ignore")
    except Exception as e:
        logger.warning("txt read failed %s: %s", path, e)
        return ""


def _extract_pdf(path: Path) -> str:
    try:
        import pypdf2 as PyPDF2  # noqa: N813
        text_parts = []
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                try:
                    text_parts.append(page.extract_text() or "")
                except Exception:
                    pass
        return "\n".join(text_parts)
    except ImportError:
        pass
    try:
        import pypdfium2 as pdfium
        doc = pdfium.PdfDocument(str(path))
        parts = []
        for page in doc:
            tp = page.get_textpage()
            parts.append(tp.get_text_range())
        return "\n".join(parts)
    except Exception as e:
        logger.warning("pdf read failed %s: %s", path, e)
        return ""


def _extract_docx(path: Path) -> str:
    try:
        from docx import Document
        doc = Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception as e:
        logger.warning("docx read failed %s: %s", path, e)
        return ""


def extract_text(path: Path) -> str:
    """Extract plain text from a file based on its extension."""
    ext = path.suffix.lower()
    if ext == ".pdf":
        return _extract_pdf(path)
    if ext == ".docx":
        return _extract_docx(path)
    return _extract_txt(path)  # .txt, .log, .csv, .md, fallback


# ── chunking ──────────────────────────────────────────────────────────────────

def _chunk_text(text: str, chunk_size: int = 400, overlap: int = 80) -> list[str]:
    """Split text into overlapping windows of roughly chunk_size words."""
    words = text.split()
    if not words:
        return []
    chunks = []
    start = 0
    while start < len(words):
        end = min(start + chunk_size, len(words))
        chunks.append(" ".join(words[start:end]))
        if end == len(words):
            break
        start += chunk_size - overlap
    return chunks


def _chunks_to_pairs(chunks: list[str]) -> list[dict]:
    """Convert consecutive chunks into prompt/completion pairs."""
    pairs = []
    for i in range(len(chunks) - 1):
        pairs.append({"prompt": chunks[i].strip(), "completion": chunks[i + 1].strip()})
    # Also produce a self-supervised 'summarise this' pair
    if chunks:
        pairs.append({
            "prompt": f"Summarise the following:\n{chunks[0]}",
            "completion": chunks[0][:300] if len(chunks[0]) > 300 else chunks[0],
        })
    return pairs


# ── public API ────────────────────────────────────────────────────────────────

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".log", ".csv", ".md", ".wav", ".mp3", ".m4a", ".ogg", ".flac", ".mp4", ".avi", ".mov", ".mkv"}


def prepare_inbox(inbox_dir: Path, batches_dir: Path, chunk_size: int = 400) -> dict:
    """
    Process all files in inbox_dir, write a JSONL batch to batches_dir,
    return a summary dict.
    """
    inbox_dir.mkdir(parents=True, exist_ok=True)
    batches_dir.mkdir(parents=True, exist_ok=True)

    files = [f for f in inbox_dir.iterdir() if f.is_file() and f.suffix.lower() in SUPPORTED_EXTENSIONS]
    if not files:
        logger.info("Inbox is empty – nothing to prepare.")
        return {"files": 0, "pairs": 0, "batch_path": None}

    all_pairs: list[dict] = []
    processed: list[str] = []
    errors: list[str] = []

    for f in files:
        try:
            text = extract_text(f)
            if not text.strip():
                errors.append(f.name)
                continue
            chunks = _chunk_text(text, chunk_size=chunk_size)
            pairs = _chunks_to_pairs(chunks)
            all_pairs.extend(pairs)
            processed.append(f.name)
            logger.info("Prepared %s: %d pairs", f.name, len(pairs))
        except Exception as e:
            logger.error("Failed to prepare %s: %s", f.name, e)
            errors.append(f.name)

    if not all_pairs:
        return {"files": len(processed), "pairs": 0, "batch_path": None, "errors": errors}

    ts = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    batch_path = batches_dir / f"batch_{ts}.jsonl"
    with open(batch_path, "w", encoding="utf-8") as out:
        for pair in all_pairs:
            out.write(json.dumps(pair, ensure_ascii=False) + "\n")

    logger.info("Wrote batch %s with %d pairs from %d files", batch_path.name, len(all_pairs), len(processed))
    return {
        "files": len(processed),
        "pairs": len(all_pairs),
        "batch_path": str(batch_path),
        "processed": processed,
        "errors": errors,
    }


def list_batches(batches_dir: Path) -> list[dict]:
    """List all JSONL batches with metadata."""
    if not batches_dir.exists():
        return []
    result = []
    for f in sorted(batches_dir.glob("*.jsonl"), reverse=True):
        try:
            lines = f.read_text(encoding="utf-8").splitlines()
            result.append({
                "name": f.name,
                "path": str(f),
                "pairs": len(lines),
                "size_kb": round(f.stat().st_size / 1024, 1),
            })
        except Exception:
            pass
    return result
